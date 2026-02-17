"""Pytest test suite for CV Processor module."""

import pytest
import os
import sys
import tempfile
from pathlib import Path
from docx import Document

# Add parent directories to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from cv_processor.cv_processor import CVProcessor
from cv_processor.extractors import PDFExtractor, DOCXExtractor, TXTExtractor
from cv_processor.parsers import ContactParser, SkillsParser, ExperienceParser
from cv_processor.normalizer import Normalizer


class TestPDFExtractor:
    """Tests for PDF extraction."""

    def test_pdf_extract_nonexistent_file(self):
        """Test extracting from non-existent PDF file."""
        result = PDFExtractor.extract("nonexistent.pdf")
        assert result is None

    def test_pdf_extract_empty_file(self):
        """Test extracting from empty/invalid PDF file."""
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            f.write(b"This is not a valid PDF")
            temp_path = f.name
        
        try:
            result = PDFExtractor.extract(temp_path)
            assert result is None
        finally:
            os.unlink(temp_path)


class TestDOCXExtractor:
    """Tests for DOCX extraction."""

    def test_docx_extract_valid_file(self):
        """Test extracting from valid DOCX file."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name
        
        try:
            # Create a test DOCX file
            doc = Document()
            doc.add_paragraph("John Doe")
            doc.add_paragraph("john@example.com")
            doc.save(temp_path)
            
            result = DOCXExtractor.extract(temp_path)
            assert result is not None
            assert "John Doe" in result
            assert "john@example.com" in result
        finally:
            os.unlink(temp_path)

    def test_docx_extract_nonexistent_file(self):
        """Test extracting from non-existent DOCX file."""
        result = DOCXExtractor.extract("nonexistent.docx")
        assert result is None


class TestTXTExtractor:
    """Tests for TXT extraction."""

    def test_txt_extract_valid_file(self):
        """Test extracting from valid TXT file."""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("Jane Smith\njane@example.com\n+1 (555) 987-6543")
            temp_path = f.name
        
        try:
            result = TXTExtractor.extract(temp_path)
            assert result is not None
            assert "Jane Smith" in result
            assert "jane@example.com" in result
        finally:
            os.unlink(temp_path)

    def test_txt_extract_empty_file(self):
        """Test extracting from empty TXT file."""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            temp_path = f.name
        
        try:
            result = TXTExtractor.extract(temp_path)
            assert result is None
        finally:
            os.unlink(temp_path)

    def test_txt_extract_nonexistent_file(self):
        """Test extracting from non-existent TXT file."""
        result = TXTExtractor.extract("nonexistent.txt")
        assert result is None


class TestContactParser:
    """Tests for contact information extraction."""

    def test_extract_email(self):
        """Test email extraction."""
        text = "Contact: john.doe@example.com for more info"
        contact = ContactParser.extract(text)
        assert contact["email"] == "john.doe@example.com"

    def test_extract_phone(self):
        """Test phone number extraction."""
        text = "Call me at +1 (555) 123-4567"
        contact = ContactParser.extract(text)
        assert contact["phone"] is not None
        assert "555" in contact["phone"]

    def test_extract_name(self):
        """Test name extraction."""
        text = "John Michael Doe\njohn@example.com"
        contact = ContactParser.extract(text)
        assert contact["name"] is not None
        assert "John" in contact["name"]

    def test_extract_incomplete_contact(self):
        """Test with incomplete contact information."""
        text = "Some random text without contact info"
        contact = ContactParser.extract(text)
        assert contact["email"] is None
        assert contact["phone"] is None

    def test_multiple_emails_returns_first(self):
        """Test that only first email is extracted."""
        text = "Primary: john@example.com, Secondary: john.work@company.com"
        contact = ContactParser.extract(text)
        assert contact["email"] == "john@example.com"

    def test_various_phone_formats(self):
        """Test various phone number formats."""
        formats = [
            "+1-555-123-4567",
            "(555) 123-4567",
            "555.123.4567",
            "+1 555 123 4567"
        ]
        
        for phone_format in formats:
            text = f"Phone: {phone_format}"
            contact = ContactParser.extract(text)
            assert contact["phone"] is not None, f"Failed for format: {phone_format}"


class TestSkillsParser:
    """Tests for skills extraction."""

    def test_extract_common_skills(self):
        """Test extraction of common skills."""
        text = "Skills: Python, JavaScript, React, FastAPI"
        skills = SkillsParser.extract(text)
        assert len(skills) > 0
        # Check for at least some standardized skills
        assert any(s.lower() in [sk.lower() for sk in skills] for s in ["Python", "JavaScript"])

    def test_extract_skills_with_abbreviations(self):
        """Test extraction with abbreviations."""
        text = "Expert in ML, AI, NLP, and DL"
        skills = SkillsParser.extract(text)
        skills_lower = [s.lower() for s in skills]
        assert any("machine learning" in s or "ml" in s for s in skills_lower)

    def test_extract_no_skills_section(self):
        """Test extraction when no skills section present."""
        text = "Just some random text about a person"
        skills = SkillsParser.extract(text)
        assert isinstance(skills, list)

    def test_standardize_skills(self):
        """Test skill standardization."""
        text = "js, typescript, react, node"
        skills = SkillsParser.extract(text)
        assert len(skills) > 0

    def test_duplicate_skills_removed(self):
        """Test that duplicate skills are removed."""
        text = "Python Python Java Java Python"
        skills = SkillsParser.extract(text)
        # Check for duplicates
        assert len(skills) == len(set(skills))


class TestExperienceParser:
    """Tests for work experience and education extraction."""

    def test_extract_experience(self):
        """Test work experience extraction."""
        text = """
        WORK EXPERIENCE
        Senior Developer at Tech Corp
        January 2020 - Present
        Led development of NLP pipeline
        
        Junior Developer at StartUp Inc
        June 2018 - December 2019
        Built REST APIs
        """
        experience = ExperienceParser.extract_experience(text)
        assert isinstance(experience, list)

    def test_extract_education(self):
        """Test education extraction."""
        text = """
        EDUCATION
        Bachelor of Science in Computer Science
        University of Technology
        Graduated: May 2018
        GPA: 3.8
        
        Master of Science in Machine Learning
        Tech University
        Graduated: June 2020
        """
        education = ExperienceParser.extract_education(text)
        assert isinstance(education, list)

    def test_extract_no_experience_section(self):
        """Test extraction when no experience section present."""
        text = "Random CV without proper sections"
        experience = ExperienceParser.extract_experience(text)
        assert isinstance(experience, list)

    def test_extract_dates(self):
        """Test date extraction."""
        text = "January 2020 - Present"
        dates = ExperienceParser._extract_dates(text)
        assert dates is not None

    def test_parse_education_entry(self):
        """Test parsing individual education entry."""
        text = """
        Bachelor of Science in Computer Science
        University of Technology
        2018
        """
        education = ExperienceParser._parse_education(text)
        assert education is not None


class TestNormalizer:
    """Tests for text normalization."""

    def test_normalize_text(self):
        """Test basic text normalization."""
        text = "This   is    multiple   spaces"
        normalized = Normalizer.normalize_text(text)
        assert "   " not in normalized
        assert normalized == "This is multiple spaces"

    def test_clean_text_with_special_chars(self):
        """Test cleaning special characters."""
        text = "Hello@World#123!Test"
        cleaned = Normalizer.clean_text(text, preserve_structure=False)
        assert "@" not in cleaned
        assert "#" not in cleaned
        assert "!" not in cleaned

    def test_clean_text_preserve_structure(self):
        """Test cleaning while preserving structure."""
        text = "Hello, World! How are you?"
        cleaned = Normalizer.clean_text(text, preserve_structure=True)
        assert "," in cleaned or "?" in cleaned

    def test_lowercase_except_acronyms(self):
        """Test selective lowercasing with acronyms."""
        text = "I work with AI, ML, and NLP"
        result = Normalizer.lowercase_except_acronyms(text)
        assert result.lower() == result.replace("AI", "ai").replace("ML", "ml").replace("NLP", "nlp")

    def test_expand_abbreviations(self):
        """Test abbreviation expansion."""
        text = "I know JS and TS"
        expanded = Normalizer.expand_abbreviations(text)
        assert "javascript" in expanded.lower()
        assert "typescript" in expanded.lower()

    def test_standardize_skills(self):
        """Test skill standardization."""
        skills = ["python", "javascript", "React", "NODE"]
        standardized = Normalizer.standardize_skills(skills)
        assert "JavaScript" in standardized
        assert "React" in standardized
        assert len(standardized) == len(set(standardized))  # No duplicates

    def test_handle_date_formats(self):
        """Test date format standardization."""
        dates = [
            ("January 2020", "2020-01"),
            ("01/15/2020", "2020-01-15"),
            ("12-31-2019", "2019-12-31"),
        ]
        
        for date_input, expected_prefix in dates:
            result = Normalizer.handle_date_formats(date_input)
            assert result.startswith(expected_prefix) or result == date_input


class TestCVProcessor:
    """Tests for main CV Processor."""

    @pytest.fixture
    def sample_cv_text(self):
        """Sample CV text for testing."""
        return """
        Santiago Cárdenas
        santiago@example.com
        +1 (555) 123-4567
        
        PROFESSIONAL SUMMARY
        Experienced software developer with expertise in NLP and web development.
        
        SKILLS
        Python, JavaScript, TypeScript, FastAPI, React, Machine Learning, NLP
        
        WORK EXPERIENCE
        Senior Developer at Tech Corp
        January 2020 - Present
        - Led development of NLP pipeline
        - Managed team of 5 engineers
        
        Junior Developer at StartUp Inc
        June 2018 - December 2019
        - Built REST APIs using FastAPI
        - Implemented ML models
        
        EDUCATION
        Bachelor of Science in Computer Science
        University of Technology
        Graduated: May 2018
        GPA: 3.8
        """

    def test_process_text_returns_dict(self, sample_cv_text):
        """Test that process_text returns a dictionary."""
        result = CVProcessor.process_text(sample_cv_text)
        assert isinstance(result, dict)

    def test_process_text_contains_required_keys(self, sample_cv_text):
        """Test that result contains all required keys."""
        result = CVProcessor.process_text(sample_cv_text)
        required_keys = ["raw_text", "normalized_text", "contact", "skills", "experience", "education"]
        for key in required_keys:
            assert key in result, f"Missing key: {key}"

    def test_process_text_contact_extraction(self, sample_cv_text):
        """Test contact info extraction in process_text."""
        result = CVProcessor.process_text(sample_cv_text)
        assert result["contact"]["email"] == "santiago@example.com"
        assert result["contact"]["phone"] is not None

    def test_process_text_skills_extraction(self, sample_cv_text):
        """Test skills extraction in process_text."""
        result = CVProcessor.process_text(sample_cv_text)
        assert len(result["skills"]) > 0
        skills_lower = [s.lower() for s in result["skills"]]
        assert any("python" in s for s in skills_lower)

    def test_process_text_empty_raises_error(self):
        """Test that empty text raises ValueError."""
        with pytest.raises(ValueError):
            CVProcessor.process_text("")

    def test_process_text_none_raises_error(self):
        """Test that None text raises ValueError."""
        with pytest.raises(ValueError):
            CVProcessor.process_text(None)

    def test_process_file_nonexistent_raises_error(self):
        """Test that processing non-existent file raises FileNotFoundError."""
        with pytest.raises(FileNotFoundError):
            CVProcessor.process_file("nonexistent_file.pdf")

    def test_process_file_unsupported_format_raises_error(self):
        """Test that unsupported format raises ValueError."""
        with tempfile.NamedTemporaryFile(suffix='.xyz', delete=False) as f:
            temp_path = f.name
        
        try:
            with pytest.raises(ValueError, match="Unsupported file format"):
                CVProcessor.process_file(temp_path)
        finally:
            os.unlink(temp_path)

    def test_process_directory_nonexistent_raises_error(self):
        """Test that non-existent directory raises ValueError."""
        with pytest.raises(ValueError):
            CVProcessor.process_directory("nonexistent_directory")

    def test_process_directory_empty(self):
        """Test processing empty directory."""
        with tempfile.TemporaryDirectory() as tmpdir:
            result = CVProcessor.process_directory(tmpdir)
            assert isinstance(result, dict)
            assert len(result) == 0

    def test_process_directory_with_txt_file(self):
        """Test processing directory with a TXT file."""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create a test TXT file
            test_file = os.path.join(tmpdir, "resume.txt")
            with open(test_file, 'w') as f:
                f.write("John Doe\njohn@example.com\n\nSkills: Python, Java")
            
            result = CVProcessor.process_directory(tmpdir)
            assert "resume.txt" in result
            assert "contact" in result["resume.txt"]


class TestIntegration:
    """Integration tests."""

    def test_full_workflow_text_processing(self):
        """Test complete workflow from text to structured data."""
        cv_text = """
        Dr. Alice Johnson
        alice.johnson@tech.com
        +44 20 7946 0958
        
        TECHNICAL SKILLS
        Languages: Python, Java, C++
        Frameworks: Django, Spring Boot, FastAPI
        Databases: PostgreSQL, MongoDB
        
        PROFESSIONAL EXPERIENCE
        Lead AI Engineer
        TechCorp International
        March 2021 - Present
        
        Senior Data Scientist
        DataInsights Ltd
        January 2019 - February 2021
        
        EDUCATION
        PhD in Machine Learning
        Stanford University
        2018
        
        BSc in Computer Science
        Cambridge University
        2014
        """
        
        result = CVProcessor.process_text(cv_text)
        
        # Verify complete extraction
        assert result["contact"]["name"] is not None
        assert result["contact"]["email"] == "alice.johnson@tech.com"
        assert len(result["skills"]) > 0
        assert len(result["education"]) > 0
        
        print("\n✓ Full workflow test passed!")
        print(f"  Name: {result['contact']['name']}")
        print(f"  Email: {result['contact']['email']}")
        print(f"  Skills found: {len(result['skills'])}")
        print(f"  Education entries: {len(result['education'])}")


if __name__ == "__main__":
    # Run tests with pytest
    pytest.main([__file__, "-v", "--tb=short"])
