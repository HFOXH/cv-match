"""DOCX text extraction module."""

from docx import Document
from typing import Optional


class DOCXExtractor:
    """Extract text from DOCX (Word) files."""

    @staticmethod
    def extract(file_path: str) -> Optional[str]:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text or None if extraction fails
        """
        try:
            doc = Document(file_path)
            text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text.append(" | ".join(row_text))
            
            return "\n".join(text) if text else None
        except Exception as e:
            print(f"Error extracting DOCX text: {str(e)}")
            return None
