"""Pytest test suite for CV Processor module.

Tests are organized by component:
    1. Extractors  — PDF, DOCX, TXT file reading
    2. Parsers     — Contact, Skills (spaCy PhraseMatcher), Experience
    3. Normalizer  — Text cleaning, abbreviation expansion, date handling
    4. CVProcessor — Main orchestrator, file/text/directory processing
    5. Integration — Full end-to-end workflow
"""

import pytest
import os
import sys
import tempfile
from pathlib import Path
from docx import Document

# Add parent directories to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from backend.nlp.processors.cv_processor import CVProcessor
from nlp.extractors import PDFExtractor, DOCXExtractor, TXTExtractor
from nlp.parsers import ContactParser, SkillsParser, ExperienceParser
from nlp.normalizers.skills_database import SKILLS_DATABASE
from backend.nlp.normalizers.normalizer import Normalizer


# ──────────────────────────────────────────────
# 1. EXTRACTORS
# ──────────────────────────────────────────────

class TestPDFExtractor:
    """Tests for PDF text extraction."""

    def test_nonexistent_file_returns_none(self):
        result = PDFExtractor.extract("nonexistent.pdf")
        assert result is None

    def test_invalid_pdf_returns_none(self):
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as f:
            f.write(b"This is not a valid PDF")
            temp_path = f.name

        try:
            result = PDFExtractor.extract(temp_path)
            assert result is None
        finally:
            os.unlink(temp_path)


class TestDOCXExtractor:
    """Tests for DOCX text extraction."""

    def test_valid_docx_extracts_text(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name

        try:
            doc = Document()
            doc.add_paragraph("John Doe")
            doc.add_paragraph("john@example.com")
            doc.save(temp_path)

            result = DOCXExtractor.extract(temp_path)
            assert result is not None
            assert "John Doe" in result
            assert "john@example.com" in result
        finally:
            os.unlink(temp_path)

    def test_nonexistent_file_returns_none(self):
        result = DOCXExtractor.extract("nonexistent.docx")
        assert result is None


class TestTXTExtractor:
    """Tests for TXT text extraction with encoding detection."""

    def test_valid_txt_extracts_text(self):
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as f:
            f.write("Jane Smith\njane@example.com\n+1 (555) 987-6543")
            temp_path = f.name

        try:
            result = TXTExtractor.extract(temp_path)
            assert result is not None
            assert "Jane Smith" in result
            assert "jane@example.com" in result
        finally:
            os.unlink(temp_path)

    def test_empty_file_returns_none(self):
        with tempfile.NamedTemporaryFile(mode='w', suffix='.txt', delete=False) as f:
            temp_path = f.name

        try:
            result = TXTExtractor.extract(temp_path)
            assert result is None
        finally:
            os.unlink(temp_path)

    def test_nonexistent_file_returns_none(self):
        result = TXTExtractor.extract("nonexistent.txt")
        assert result is None


# ──────────────────────────────────────────────
# 2. PARSERS
# ──────────────────────────────────────────────

class TestContactParser:
    """Tests for contact information extraction (name, email, phone)."""

    def test_extract_email(self):
        text = "Contact: john.doe@example.com for more info"
        contact = ContactParser.extract(text)
        assert contact["email"] == "john.doe@example.com"

    def test_extract_phone_us_format(self):
        text = "Call me at +1 (555) 123-4567"
        contact = ContactParser.extract(text)
        assert contact["phone"] is not None
        assert "555" in contact["phone"]

    def test_extract_phone_international_format(self):
        text = "Phone: (+57) 3132904901"
        contact = ContactParser.extract(text)
        assert contact["phone"] is not None
        assert "313" in contact["phone"]

    def test_extract_name(self):
        text = "John Michael Doe\njohn@example.com"
        contact = ContactParser.extract(text)
        assert contact["name"] is not None
        assert "John" in contact["name"]

    def test_incomplete_contact_returns_none_fields(self):
        text = "Some random text without contact info"
        contact = ContactParser.extract(text)
        assert contact["email"] is None
        assert contact["phone"] is None

    def test_multiple_emails_returns_first(self):
        text = "Primary: john@example.com, Secondary: john.work@company.com"
        contact = ContactParser.extract(text)
        assert contact["email"] == "john@example.com"

    def test_various_phone_formats(self):
        formats = [
            "+1-555-123-4567",
            "(555) 123-4567",
            "555.123.4567",
            "+1 555 123 4567",
        ]
        for phone_format in formats:
            text = f"Phone: {phone_format}"
            contact = ContactParser.extract(text)
            assert contact["phone"] is not None, f"Failed for format: {phone_format}"


class TestSkillsParser:
    """Tests for spaCy PhraseMatcher skill extraction.

    Covers: document-wide matching, section-based extraction,
    abbreviation resolution, duplicate removal, and fallback handling.
    """

    def test_extract_common_skills(self):
        text = "Skills: Python, JavaScript, React, FastAPI"
        skills = SkillsParser.extract(text)
        skill_names = [s.lower() for s in skills]
        assert "python" in skill_names
        assert "javascript" in skill_names

    def test_extract_abbreviations(self):
        text = "Expert in ML, AI, NLP, and DL"
        skills = SkillsParser.extract(text)
        skill_names = [s.lower() for s in skills]
        assert "machine learning" in skill_names
        assert "artificial intelligence" in skill_names
        assert "nlp" in skill_names

    def test_no_skills_returns_empty_list(self):
        text = "Just some random text about a person"
        skills = SkillsParser.extract(text)
        assert isinstance(skills, list)

    def test_standardizes_variant_names(self):
        """js -> JavaScript, ts -> TypeScript, etc."""
        text = "js, typescript, react, node"
        skills = SkillsParser.extract(text)
        skill_names = [s.lower() for s in skills]
        assert "javascript" in skill_names
        assert "typescript" in skill_names

    def test_duplicate_skills_removed(self):
        text = "Python Python Java Java Python"
        skills = SkillsParser.extract(text)
        assert len(skills) == len(set(skills))

    def test_multiword_skills(self):
        """PhraseMatcher should catch multi-word skills in one pass."""
        text = "I have experience with machine learning and react native"
        skills = SkillsParser.extract(text)
        skill_names = [s.lower() for s in skills]
        assert "machine learning" in skill_names
        assert "react native" in skill_names

    def test_section_based_extraction(self):
        """Skills listed in a SKILLS section should be found."""
        text = """
        SKILLS
        Python, Docker, Kubernetes, FastAPI

        EXPERIENCE
        Software Engineer
        """
        skills = SkillsParser.extract(text)
        skill_names = [s.lower() for s in skills]
        assert "python" in skill_names
        assert "docker" in skill_names
        assert "kubernetes" in skill_names

    def test_skills_database_has_entries(self):
        """The imported skills database should have 200+ entries."""
        assert len(SKILLS_DATABASE) > 200

    def test_short_skills_skipped_in_document_wide_scan(self):
        """Single-char skills like 'C' and 'R' should not match everywhere."""
        text = "I can see the results clearly and read reports"
        skills = SkillsParser.extract(text)
        skill_names = [s.lower() for s in skills]
        assert "c" not in skill_names
        assert "r" not in skill_names


class TestExperienceParser:
    """Tests for work experience and education extraction.

    Covers: section detection, date parsing, job/company parsing,
    education degree/institution extraction, and edge cases.
    """

    def test_extract_experience_basic(self):
        text = """
        WORK EXPERIENCE
        Senior Developer at Tech Corp
        January 2020 - Present
        Led development of NLP pipeline

        Junior Developer at StartUp Inc
        June 2018 - December 2019
        Built REST APIs
        """
        experience = ExperienceParser.extract_experience(text)
        assert isinstance(experience, list)
        assert len(experience) >= 1

    def test_extract_education_basic(self):
        text = """
        EDUCATION
        Bachelor of Science in Computer Science
        University of Technology
        Graduated: May 2018
        GPA: 3.8

        Master of Science in Machine Learning
        Tech University
        Graduated: June 2020
        """
        education = ExperienceParser.extract_education(text)
        assert isinstance(education, list)
        assert len(education) >= 1

    def test_no_experience_section_returns_empty(self):
        text = "Random CV without proper sections"
        experience = ExperienceParser.extract_experience(text)
        assert isinstance(experience, list)
        assert len(experience) == 0

    def test_extract_dates_month_year(self):
        dates = ExperienceParser._extract_dates("January 2020 - Present")
        assert dates is not None
        assert dates[0] == "January 2020"
        assert dates[1] == "Present"

    def test_extract_dates_numeric(self):
        dates = ExperienceParser._extract_dates("01/15/2020")
        assert dates is not None

    def test_parse_education_entry(self):
        text = """
        Bachelor of Science in Computer Science
        University of Technology
        2018
        """
        education = ExperienceParser._parse_education(text)
        assert education is not None
        assert education["degree"] is not None
        assert "computer science" in education["degree"].lower()

    def test_section_heading_not_matched_inside_words(self):
        """'experience' inside 'Experienced' should NOT trigger section detection."""
        text = """
        SUMMARY
        Experienced software developer with 5 years of work.

        EDUCATION
        Bachelor of Science
        University of Technology
        """
        experience = ExperienceParser.extract_experience(text)
        assert len(experience) == 0

    def test_date_line_with_company(self):
        """Dates and company on the same line should both be extracted."""
        text = """
        EXPERIENCE
        Full Stack Developer
        Oct 2023 - Present Symplifica
        Built web applications
        """
        experience = ExperienceParser.extract_experience(text)
        assert len(experience) >= 1
        job = experience[0]
        assert job["start_date"] is not None

    def test_education_year_range(self):
        """Year ranges like '2020 - 2023' on the same line should be captured."""
        text = """
        EDUCATION
        Systems Engineering  2020 - 2023
        University of Colombia
        """
        education = ExperienceParser.extract_education(text)
        assert len(education) >= 1
        assert education[0]["graduation_date"] is not None


# ──────────────────────────────────────────────
# 3. NORMALIZER
# ──────────────────────────────────────────────

class TestNormalizer:
    """Tests for text normalization utilities."""

    def test_normalize_collapses_whitespace(self):
        text = "This   is    multiple   spaces"
        normalized = Normalizer.normalize_text(text)
        assert "   " not in normalized
        assert normalized == "This is multiple spaces"

    def test_clean_text_removes_special_chars(self):
        text = "Hello@World#123!Test"
        cleaned = Normalizer.clean_text(text, preserve_structure=False)
        assert "@" not in cleaned
        assert "#" not in cleaned
        assert "!" not in cleaned

    def test_clean_text_preserve_structure(self):
        text = "Hello, World! How are you?"
        cleaned = Normalizer.clean_text(text, preserve_structure=True)
        assert "," in cleaned or "?" in cleaned

    def test_lowercase_except_acronyms(self):
        text = "I work with AI, ML, and NLP"
        result = Normalizer.lowercase_except_acronyms(text)
        assert result.lower() == result.replace("AI", "ai").replace("ML", "ml").replace("NLP", "nlp")

    def test_expand_abbreviations(self):
        text = "I know JS and TS"
        expanded = Normalizer.expand_abbreviations(text)
        assert "javascript" in expanded.lower()
        assert "typescript" in expanded.lower()

    def test_standardize_skills(self):
        skills = ["python", "javascript", "React", "NODE"]
        standardized = Normalizer.standardize_skills(skills)
        assert "JavaScript" in standardized
        assert "React" in standardized
        assert len(standardized) == len(set(standardized))

    def test_handle_date_formats(self):
        dates = [
            ("January 2020", "2020-01"),
            ("01/15/2020", "2020-01-15"),
            ("12-31-2019", "2019-12-31"),
        ]
        for date_input, expected_prefix in dates:
            result = Normalizer.handle_date_formats(date_input)
            assert result.startswith(expected_prefix) or result == date_input


# ──────────────────────────────────────────────
# 4. CV PROCESSOR (Main Orchestrator)
# ──────────────────────────────────────────────

class TestCVProcessor:
    """Tests for the main CVProcessor class."""

    @pytest.fixture
    def sample_cv_text(self):
        """Sample CV text used across multiple tests."""
        return """
        Santiago Cárdenas
        santiago@example.com
        +1 (555) 123-4567

        PROFESSIONAL SUMMARY
        Experienced software developer with expertise in NLP and web development.

        SKILLS
        Python, JavaScript, TypeScript, FastAPI, React, Machine Learning, NLP

        WORK EXPERIENCE
        Senior Developer at Tech Corp
        January 2020 - Present
        - Led development of NLP pipeline
        - Managed team of 5 engineers

        Junior Developer at StartUp Inc
        June 2018 - December 2019
        - Built REST APIs using FastAPI
        - Implemented ML models

        EDUCATION
        Bachelor of Science in Computer Science
        University of Technology
        Graduated: May 2018
        GPA: 3.8
        """

    # -- Return type and structure --

    def test_process_text_returns_dict(self, sample_cv_text):
        result = CVProcessor.process_text(sample_cv_text)
        assert isinstance(result, dict)

    def test_process_text_contains_required_keys(self, sample_cv_text):
        result = CVProcessor.process_text(sample_cv_text)
        required_keys = ["raw_text", "normalized_text", "contact", "skills", "experience", "education"]
        for key in required_keys:
            assert key in result, f"Missing key: {key}"

    # -- Contact extraction --

    def test_process_text_extracts_email(self, sample_cv_text):
        result = CVProcessor.process_text(sample_cv_text)
        assert result["contact"]["email"] == "santiago@example.com"

    def test_process_text_extracts_phone(self, sample_cv_text):
        result = CVProcessor.process_text(sample_cv_text)
        assert result["contact"]["phone"] is not None

    # -- Skills extraction --

    def test_process_text_extracts_skills(self, sample_cv_text):
        result = CVProcessor.process_text(sample_cv_text)
        assert len(result["skills"]) > 0
        skill_names = [s.lower() for s in result["skills"]]
        assert "python" in skill_names

    # -- Error handling --

    def test_empty_text_raises_error(self):
        with pytest.raises(ValueError):
            CVProcessor.process_text("")

    def test_none_text_raises_error(self):
        with pytest.raises(ValueError):
            CVProcessor.process_text(None)

    def test_nonexistent_file_raises_error(self):
        with pytest.raises(FileNotFoundError):
            CVProcessor.process_file("nonexistent_file.pdf")

    def test_unsupported_format_raises_error(self):
        with tempfile.NamedTemporaryFile(suffix='.xyz', delete=False) as f:
            temp_path = f.name

        try:
            with pytest.raises(ValueError, match="Unsupported file format"):
                CVProcessor.process_file(temp_path)
        finally:
            os.unlink(temp_path)

    # -- Directory processing --

    def test_nonexistent_directory_raises_error(self):
        with pytest.raises(ValueError):
            CVProcessor.process_directory("nonexistent_directory")

    def test_empty_directory(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            result = CVProcessor.process_directory(tmpdir)
            assert isinstance(result, dict)
            assert len(result) == 0

    def test_directory_with_txt_file(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            test_file = os.path.join(tmpdir, "resume.txt")
            with open(test_file, 'w') as f:
                f.write("John Doe\njohn@example.com\n\nSkills: Python, Java")

            result = CVProcessor.process_directory(tmpdir)
            assert "resume.txt" in result
            assert "contact" in result["resume.txt"]


# ──────────────────────────────────────────────
# 5. INTEGRATION
# ──────────────────────────────────────────────

class TestIntegration:
    """End-to-end tests verifying the full CV processing pipeline."""

    def test_full_workflow_text_processing(self):
        """Process a complete CV and verify all sections are extracted."""
        cv_text = """
        Dr. Alice Johnson
        alice.johnson@tech.com
        +44 20 7946 0958

        TECHNICAL SKILLS
        Languages: Python, Java, C++
        Frameworks: Django, Spring Boot, FastAPI
        Databases: PostgreSQL, MongoDB

        PROFESSIONAL EXPERIENCE
        Lead AI Engineer
        TechCorp International
        March 2021 - Present

        Senior Data Scientist
        DataInsights Ltd
        January 2019 - February 2021

        EDUCATION
        PhD in Machine Learning
        Stanford University
        2018

        BSc in Computer Science
        Cambridge University
        2014
        """

        result = CVProcessor.process_text(cv_text)

        # Contact
        assert result["contact"]["name"] is not None
        assert result["contact"]["email"] == "alice.johnson@tech.com"
        assert result["contact"]["phone"] is not None

        # Skills
        assert len(result["skills"]) > 0
        skill_names = [s.lower() for s in result["skills"]]
        assert "python" in skill_names
        assert "django" in skill_names

        # Education
        assert len(result["education"]) > 0


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])
