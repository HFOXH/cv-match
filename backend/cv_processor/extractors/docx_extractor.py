import logging
from typing import Optional

from docx import Document

logger = logging.getLogger(__name__)


class DOCXExtractor:
    """
    Extract text from DOCX (Microsoft Word) files.

    This extractor reads all paragraphs and tables from a DOCX document and
    concatenates their text into a single string. Useful for processing resumes,
    reports, or any structured Word documents.

    Methods
    -------
    extract(file_path)
        Extracts and returns the text from the DOCX file, or None if extraction fails.
    """
    @staticmethod
    def extract(file_path: str) -> Optional[str]:
        try:
            doc = Document(file_path)
            text = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text.append(" | ".join(row_text))

            return "\n".join(text) if text else None
        except Exception as e:
            logger.error("Error extracting DOCX text: %s", e)
            return None
